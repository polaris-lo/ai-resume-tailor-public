"""
简历定制核心模块

根据岗位 JD，调用 LLM（OpenAI 兼容格式）生成针对性修改建议，输出定制版 .docx。

格式规范（自动强制执行）：
- 全文宋体五号（10.5pt），含 ascii/hAnsi/eastAsia/cs 四个字符集
- 加粗规则：只加粗关键词（技能类别名、条目类型词、关键数字），不整句加粗
- 无空段落，无嵌入 sectPr
"""

from __future__ import annotations

import json
import shutil
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from openai import OpenAI

# 将项目根目录加入 sys.path，以便 import config
sys.path.insert(0, str(Path(__file__).parent.parent.parent))
from config import load_config

OUTPUT_DIR = Path(__file__).parent.parent.parent / "output"


def _get_client():
    cfg = load_config()
    return OpenAI(
        api_key=cfg["llm"]["api_key"],
        base_url=cfg["llm"]["base_url"],
    ), cfg["llm"]["model"]


def _get_user_name() -> str:
    return load_config()["user_name"]


def _get_base_resume() -> Path:
    cfg = load_config()
    p = Path(cfg["resume_path"]).expanduser()
    return p


# ──────────────────────────────────────────────────────────────
# 读取简历内容
# ──────────────────────────────────────────────────────────────

def get_resume_content(doc_path: Path) -> list[dict]:
    """提取简历段落结构，同时兼容表格布局。

    返回两类条目：
    - in_table=False：正文段落，有 index，可被 LLM 修改
    - in_table=True ：表格单元格文字，index=-1，仅供 LLM 阅读，不可修改
    """
    doc = Document(str(doc_path))
    result = []

    for i, para in enumerate(doc.paragraphs):
        if not para.text.strip():
            continue
        result.append({
            "index": i,
            "is_header": para.style.name == "Normal",
            "text": para.text,
            "in_table": False,
        })

    # 提取表格内容（合并相邻重复单元格，避免 python-docx 的合并单元格重复读取）
    seen = set()
    for table in doc.tables:
        for row in table.rows:
            row_texts = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text and cell_text not in seen:
                    seen.add(cell_text)
                    row_texts.append(cell_text)
            if row_texts:
                result.append({
                    "index": -1,
                    "is_header": False,
                    "text": " | ".join(row_texts),
                    "in_table": True,
                })

    return result


# ──────────────────────────────────────────────────────────────
# 调用 LLM
# ──────────────────────────────────────────────────────────────

_SYSTEM_PROMPT = """\
你是一位专业的简历顾问，擅长根据岗位要求对简历进行针对性修改。

【格式规范（必须严格遵守）】
1. 只修改"专业技能"和"实习经历"中的描述句，不修改基本信息、教育背景、节标题、公司/时间行
2. 每个修改段落用 segments 表示，segments 是文本片段数组，每片段有 text 和 bold 两个字段
3. bold=true 仅用于：技能类别名、条目类型词、关键数字/比例，绝不整句加粗
4. 修改要简洁量化，突出与 JD 匹配的能力，风格与原简历一致

【特别注意】
- 简历中标注「表格内容」的部分仅供参考，不得出现在返回 JSON 中
- JSON 中只能使用「正文段落」区域里列出的 index 编号

【返回格式】
只返回一个 JSON 数组，不要任何其他文字：
[
  {
    "para_index": <段落在原文中的 index 编号>,
    "reason": "<一句话说明修改原因>",
    "segments": [
      {"text": "关键词", "bold": true},
      {"text": "：描述文字", "bold": false}
    ]
  }
]
"""


def _parse_llm_json(raw: str) -> list[dict]:
    """解析 LLM 返回的 JSON，兼容 ```json ... ``` 包裹。"""
    raw = raw.strip()
    if "```" in raw:
        for part in raw.split("```"):
            part = part.strip().lstrip("json").strip()
            if part.startswith("["):
                raw = part
                break
    return json.loads(raw)


def _build_resume_str(resume_paras: list[dict]) -> str:
    normal_lines = []
    table_lines = []
    for p in resume_paras:
        if p.get("in_table"):
            table_lines.append(f"  {p['text']}")
        else:
            tag = "[节标题]" if p["is_header"] else "       "
            normal_lines.append(f"  {p['index']:2d}. {tag} {p['text']}")

    parts = ["【正文段落（可修改，使用上方 index 编号）】"] + normal_lines
    if table_lines:
        parts += ["", "【表格内容（仅供参考，不可修改，不得出现在 JSON 中）】"] + table_lines
    return "\n".join(parts)


def call_llm(jd: str, resume_paras: list[dict]) -> list[dict]:
    """调用 LLM，返回段落修改列表。"""
    client, model = _get_client()
    user_msg = f"【岗位描述】\n{jd}\n\n【当前简历段落】\n{_build_resume_str(resume_paras)}"
    response = client.chat.completions.create(
        model=model,
        max_tokens=4096,
        messages=[
            {"role": "system", "content": _SYSTEM_PROMPT},
            {"role": "user", "content": user_msg},
        ],
    )
    return _parse_llm_json(response.choices[0].message.content)


def call_llm_with_feedback(
    jd: str,
    resume_paras: list[dict],
    prev_modifications: list[dict],
    feedback: str,
) -> list[dict]:
    """基于用户反馈，重新调用 LLM 生成修改建议。"""
    client, model = _get_client()
    orig_map = {p["index"]: p["text"] for p in resume_paras}
    prev_lines = []
    for m in prev_modifications:
        orig = orig_map.get(m["para_index"], "")
        new_text = "".join(s["text"] for s in m.get("segments", []))
        prev_lines.append(
            f"  段落 {m['para_index']}: {orig!r} → {new_text!r}  ({m.get('reason', '')})"
        )
    user_msg = (
        f"【岗位描述】\n{jd}\n\n"
        f"【当前简历段落】\n{_build_resume_str(resume_paras)}\n\n"
        f"【上一版修改建议】\n" + "\n".join(prev_lines) + "\n\n"
        f"【用户反馈】\n{feedback}\n\n"
        "请根据用户反馈调整修改建议，重新输出完整 JSON 数组。"
    )
    response = client.chat.completions.create(
        model=model,
        max_tokens=4096,
        messages=[
            {"role": "system", "content": _SYSTEM_PROMPT},
            {"role": "user", "content": user_msg},
        ],
    )
    return _parse_llm_json(response.choices[0].message.content)


def suggest_job_name(jd: str) -> str:
    """根据 JD 内容自动推断岗位名称（不超过10个汉字）。"""
    client, model = _get_client()
    response = client.chat.completions.create(
        model=model,
        max_tokens=20,
        messages=[{
            "role": "user",
            "content": (
                "根据以下岗位描述，给出一个简短的岗位名称（不超过10个汉字，"
                "如「AI产品经理」、「数据分析师」），只返回名称本身，不要任何其他文字：\n\n"
                + jd[:800]
            ),
        }],
    )
    return response.choices[0].message.content.strip()


def format_modifications(modifications: list[dict], resume_paras: list[dict]) -> str:
    """格式化修改建议，便于展示给用户确认。"""
    orig_map = {p["index"]: p["text"] for p in resume_paras}
    lines = [f"共 {len(modifications)} 处修改建议：\n"]
    for i, m in enumerate(modifications, 1):
        orig = orig_map.get(m["para_index"], "（原文未找到）")
        new_text = "".join(s["text"] for s in m.get("segments", []))
        lines.append(f"[{i}] 段落 {m['para_index']} — {m.get('reason', '')}")
        lines.append(f"    原文：{orig}")
        lines.append(f"    修改：{new_text}")
        lines.append("")
    return "\n".join(lines)


# ──────────────────────────────────────────────────────────────
# 应用修改
# ──────────────────────────────────────────────────────────────

def _set_runs(para, segments: list[dict]) -> None:
    """替换段落所有 run，统一应用宋体五号 + bold 规则。"""
    p = para._p
    for r in list(p.findall(qn("w:r"))):
        p.remove(r)
    for seg in segments:
        run = para.add_run(seg["text"])
        run.bold = seg["bold"]
        run.font.size = Pt(10.5)
        rPr = run._r.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
            rFonts.set(qn(attr), "宋体")


def apply_modifications(base_path: Path, modifications: list[dict], output_path: Path) -> None:
    """将修改应用到 docx 副本，同时清理空段落和嵌入 sectPr。"""
    shutil.copy(str(base_path), str(output_path))
    doc = Document(str(output_path))

    # 应用段落替换
    mod_map = {m["para_index"]: m for m in modifications}
    for i, para in enumerate(doc.paragraphs):
        if i in mod_map:
            _set_runs(para, mod_map[i]["segments"])

    # 清理空段落
    body = doc.element.body
    for p in list(body.findall(qn("w:p"))):
        texts = [t.text or "" for t in p.findall(".//" + qn("w:t"))]
        if not "".join(texts).strip():
            body.remove(p)

    # 清理嵌入 sectPr（连续分节符）
    for para in doc.paragraphs:
        pPr = para._p.find(qn("w:pPr"))
        if pPr is not None:
            sectPr = pPr.find(qn("w:sectPr"))
            if sectPr is not None:
                pPr.remove(sectPr)

    # 统一全文字体为宋体五号
    for para in doc.paragraphs:
        for run in para.runs:
            run.font.size = Pt(10.5)
            rPr = run._r.get_or_add_rPr()
            rFonts = rPr.get_or_add_rFonts()
            for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
                rFonts.set(qn(attr), "宋体")

    doc.save(str(output_path))


# ──────────────────────────────────────────────────────────────
# 打招呼消息
# ──────────────────────────────────────────────────────────────

_GREETING_SYSTEM_PROMPT = """\
你是一位求职文案助手，帮助求职者在 BOSS 直聘上撰写发送给 HR 的第一条自定义打招呼消息。

【写作要求】
1. 开头一句话简介自己（身份 + 专业方向），语气自然
2. 接着列出 2～4 个与该 JD 最匹配的技能点，每个技能后用括号简短注明在哪段经历中具体用过（从提供的简历内容中提取，保持真实）
3. 结尾一句表达对岗位的兴趣，邀请进一步沟通
4. 总字数控制在 100 字以内，语气真诚不浮夸，适合职场初次联系
5. 全部用中文，开头可以"您好"，不要"尊敬的HR"之类
6. 只返回消息正文，不要任何解释或标注
"""


def generate_greeting(jd: str, resume_paras: list[dict]) -> str:
    """根据 JD 和简历内容，生成 BOSS 直聘打招呼消息。"""
    client, model = _get_client()
    user_msg = f"【岗位描述】\n{jd}\n\n【简历内容】\n{_build_resume_str(resume_paras)}"
    response = client.chat.completions.create(
        model=model,
        max_tokens=300,
        messages=[
            {"role": "system", "content": _GREETING_SYSTEM_PROMPT},
            {"role": "user", "content": user_msg},
        ],
    )
    return response.choices[0].message.content.strip()


# ──────────────────────────────────────────────────────────────
# 主入口
# ──────────────────────────────────────────────────────────────

def tailor_resume(
    jd: str,
    output_name: str,
    modifications: list[dict],
    base_path: Path | None = None,
    output_dir: Path = OUTPUT_DIR,
) -> Path:
    """将已确认的修改应用到 docx，返回输出文件路径。"""
    if base_path is None:
        base_path = _get_base_resume()
    user_name = _get_user_name()
    output_dir.mkdir(exist_ok=True)
    output_path = output_dir / f"{user_name}简历_{output_name}版.docx"
    apply_modifications(base_path, modifications, output_path)
    return output_path
